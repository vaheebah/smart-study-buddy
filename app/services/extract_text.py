import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import re
from pathlib import Path

def extract_text_from_pdf(file_path: str) -> dict:
    """Extract text from PDF file"""
    try:
        text = ""
        page_count = 0
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return {"text": clean_text(text), "pages": page_count, "language": "en"}
    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(file_path: str) -> dict:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return {"text": clean_text(text), "pages": len(doc.paragraphs), "language": "en"}
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def extract_text_from_image(file_path: str) -> dict:
    """Extract text from image using OCR"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return {"text": clean_text(text), "pages": 1, "language": "en"}
    except Exception as e:
        # If OCR fails, return placeholder
        return {"text": "Image text extraction failed. Please use text documents.", "pages": 1, "language": "en"}

def extract_text_from_txt(file_path: str) -> dict:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {"text": clean_text(text), "pages": 1, "language": "en"}
    except Exception as e:
        raise Exception(f"TXT extraction failed: {str(e)}")

def clean_text(text: str) -> str:
    """Clean extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove non-printable characters
    text = ''.join(char for char in text if char.isprintable() or char == '\n')
    return text.strip()

def extract_text(file_path: str) -> dict:
    """Detect file type and extract text"""
    path = Path(file_path)
    suffix = path.suffix.lower()
    
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    elif suffix in [".txt"]:
        return extract_text_from_txt(file_path)
    elif suffix in [".png", ".jpg", ".jpeg"]:
        return extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
